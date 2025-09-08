from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_borders(cell):
    """
    Add borders to the table cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create a new element for cell borders
    borders = OxmlElement('w:tcBorders')
    for border_side in ('top', 'bottom', 'left', 'right'):
        border_element = OxmlElement(f'w:{border_side}')
        borders.append(border_element)

    # Set border properties
    for border in borders:
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border size in half-points
        border.set(qn('w:color'), 'auto')

    # Add borders to the cell
    tcPr.append(borders)

def create_practical_training_log_book(department, student_name, reg_no, company, week_no, from_date, to_date, data_dictionary):

    """
    data_dictionary = {
        monday:{
            'date': date,
            'activity': activity
        },
        tuesday:{
            'date': date,
            'activity': activity
        }
    }
    """
    # Create a new Document
    doc = Document()

    # Title
    title = doc.add_heading('NAME OF MY UNIVERSITY', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # College and Department
    college_heading = doc.add_heading('COLLEGE OF MY UNIVERSITY', level=2)
    college_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    department = doc.add_heading(f'DEPARTMENT OF {department}', level=2)
    department.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Practical Training Log - Book
    pt_title = doc.add_heading('PRACTICAL TRAINING LOG – BOOK', level=2)
    pt_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    # Student Information
    student_table = doc.add_table(rows=1, cols=2)
    student_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set column widths (50% for each column)
    student_table.columns[0].width = Inches(4)
    student_table.columns[1].width = Inches(2)

    # Add text to the cells
    cell_0_0 = student_table.cell(0, 0)
    cell_0_1 = student_table.cell(0, 1)

    cell_0_0.text = f'STUDENTS NAME: {student_name}'
    cell_0_1.text = f'REG. NO: {reg_no}'

    # add another row with a single column for company name
    company_row = student_table.add_row()
    company_cell = company_row.cells[0]
    company_cell.text = f'COMPANY/INSTITUTION: {company}'
    company_cell.merge(student_table.cell(1, 1)) 

    week_log_table = doc.add_table(rows=1, cols=3)
    week_log_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Week number, from date, and to date cells
    week_number_cell = week_log_table.cell(0, 0)
    week_number_cell.text = f'WEEK NO: {week_no}'
    week_number_cell.paragraphs[0].paragraph_format.space_before = Inches(0.1)

    from_date_cell = week_log_table.cell(0, 1)
    from_date_cell.text = f'FROM: {from_date}'
    from_date_cell.paragraphs[0].paragraph_format.space_before = Inches(0.1)

    to_date_cell = week_log_table.cell(0, 2)
    to_date_cell.text = f'TO: {to_date}'
    to_date_cell.paragraphs[0].paragraph_format.space_before = Inches(0.1)

    # set cell widths 30% each
    week_log_table.columns[0].width = Inches(1.5)
    week_log_table.columns[1].width = Inches(2.25)
    week_log_table.columns[2].width = Inches(2.25)
    
    # Add borders to the table cells
    for cell in week_log_table._cells:
        set_cell_borders(cell)

                
    # add an empty paragraph with 2 new line breaks
    doc.add_paragraph('\n')

    # Days and Activities
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']

    # Create a table for Days and Activities
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False

    # Set vertical alignment to top for all cells
    for cell in table._cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Add borders to the table cells
        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)

    # Set column widths
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4)

    # Add Days and Activities column headers
    cell_0_0 = table.cell(0, 0)
    cell_0_1 = table.cell(0, 1)

    cell_0_0.text = 'DAY / DATE'
    cell_0_1.text = 'ACTIVITY'

    # Make the text in the cells bold
    for cell in [cell_0_0, cell_0_1]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Align the text in the cells at the center
    cell_0_0.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_0_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add spacing before the headers
    cell_0_0.paragraphs[0].paragraph_format.space_before = Inches(.1)
    cell_0_1.paragraphs[0].paragraph_format.space_before = Inches(.1)

    # Populate the table with Days, Dates, and an empty cell for activities
    for i, day in enumerate(days):
        table.cell(i + 1, 0).text = f" {day} \n \n {data_dictionary[day]['date']}"
        table.cell(i + 1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.cell(i + 1, 1).text = data_dictionary[day]['activity']

    doc.add_page_break() # page break

    # Details of the Main Job of the Week
    doc.add_heading('Details Of the Main Job of the Week', level=2)

    # Operation and Machinery/Tools Used
    machinery_heading = doc.add_heading('Machinery/Tools Used', level=3)

    # Add a table for machinery/tools used
    machinery_table = doc.add_table(rows=8, cols=2)  # Add 1 row for the header
    machinery_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    machinery_table.autofit = False
    machinery_table.columns[0].width = Inches(3)
    machinery_table.columns[1].width = Inches(3)

    # Add borders to the table cells
    for cell in machinery_table._cells:
        set_cell_borders(cell)

    # Add the column headers (single-row header)
    cell_0_0 = machinery_table.cell(0, 0)
    cell_0_1 = machinery_table.cell(0, 1)

    cell_0_0.text = 'Operation: '
    cell_0_1.text = 'Machinery/Tools Used'

    # Make the text in the header cells bold
    for cell in [cell_0_0, cell_0_1]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Align the text in the header cells at the center
    cell_0_0.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_0_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Populate the rest of the table with empty cells
    for row in machinery_table.rows[1:]:
        for cell in row.cells:
            cell.text = ' '


    # table for comments from industrial supervisor, 1 row, 2 columns comments and signature
    comments_table = doc.add_table(rows=1, cols=2)
    comments_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    comments_table.columns[0].width = Inches(2)
    comments_table.columns[1].width = Inches(4)

    # Add borders to the table cells
    for cell in comments_table._cells:
        set_cell_borders(cell)

    # Add the column headers
    comments_table.cell(0, 0).text = 'Comments from Industrial Supervisor'
    comments_table.cell(0, 1).text = '\n'

    # table for Name and Signature
    name_table = doc.add_table(rows=1, cols=2)
    name_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_table.columns[0].width = Inches(3)
    name_table.columns[1].width = Inches(3)

    # Add borders to the table cells
    for cell in name_table._cells:
        set_cell_borders(cell)

    # Add the column headers
    name_table.cell(0, 0).text = '\nName: ............................................................'
    name_table.cell(0, 1).text = '\nSignature: .......................................'

    doc.add_page_break() #page break

    # Detailed Diagram of the Main Job
    doc.add_heading('Detailed Diagram of the Main Job', level=2)
    diagram_paragraph = doc.add_paragraph()
    diagram_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    # Save the Document
    filepath = f"docs/{reg_no}-week-{week_no}-practical-training-logbook.docx"
    doc.save(filepath) 

    return filepath



# Sample data
department = "COMPUTER SCIENCE"
student_name = "mgasa.loucat1@gmail.com"
reg_no = "2021-04-02100"
company = "ABC Technologies"
week_no = 1
from_date = "2023-07-01"
to_date = "2023-07-07"

data_dictionary = {
    'Monday': {
        'date': '2023-07-01',
        'activity': 'Worked on project X'
    },
    'Tuesday': {
        'date': '2023-07-02',
        'activity': 'Attended team meeting'
    },
    'Wednesday': {
        'date': '2023-07-03',
        'activity': 'Researched new technologies'   
    },
    'Thursday': {
        'date': '2023-07-04',
        'activity': 'Debugged code'
    },
    'Friday': {
        'date': '2023-07-05',
        'activity': 'Prepared project presentation'
    }
}

# Call the function to generate the logbook
filepath = create_practical_training_log_book(department, student_name, reg_no, company, week_no, from_date, to_date, data_dictionary)

print(f"Logbook generated and saved to: {filepath}")
