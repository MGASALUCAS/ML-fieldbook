import os
from main.settings import MEDIA_ROOT, BASE_DIR
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from html2docx import html2docx


def set_cell_borders(cell):
    """
    Add borders to the table cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create a new element for cell borders
    borders = OxmlElement("w:tcBorders")
    for border_side in ("top", "bottom", "left", "right"):
        border_element = OxmlElement(f"w:{border_side}")
        borders.append(border_element)

    # Set border properties
    for border in borders:
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")  # Border size in half-points
        border.set(qn("w:color"), "auto")

    # Add borders to the cell
    tcPr.append(borders)


def create_practical_training_log_book(
    department,
    student_name,
    reg_no,
    company,
    week_no,
    from_date,
    to_date,
    data_dictionary,
    operations,
    activity_diagram,
):
    """
    Creates a Word document for a student's practical training log book.

    Args:
        department (str): Name of the department.
        student_name (str): Name of the student.
        reg_no (str): Student's registration number.
        company (str): Name of the company/institution.
        week_no (str): Week number of the training.
        from_date (str): Start date of the week (e.g., '2023-10-02').
        to_date (str): End date of the week (e.g., '2023-10-06').
        data_dictionary (dict): Dictionary mapping days to date and activity.
            Format: {'Monday': {'date': str, 'activity': str}, ...}
        operations (list): List of dictionaries with operation and machinery.
            Format: [{'operation': str, 'machinery': str}, ...]
        activity_diagram (str): File path to the activity diagram image.

    Returns:
        str: Filepath of the saved Word document.

    Raises:
        ValueError: If data_dictionary is missing required days.
        FileNotFoundError: If activity_diagram file does not exist.
    """

    """
    data_dictionary = {
        monday:{
            'date': date,
            'activity': activity
        },
        tuesday:{
            'date': date,
            'activity': activity
        }
    }

    operations = [ {
            'operation': operation,
            'machinery': machinery},

            {'operation': operation,
            'machinery': machinery}
            ]

    """
    # capitalize 'department, student_name, company'
    department = department.upper()
    student_name = student_name.upper()
    company = company.upper()

    # Create a new Document
    doc = Document()

    # Title
    title = doc.add_heading("UNIVERSITY OF DAR ES SALAAM", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # College and Department
    college_heading = doc.add_heading(
        "COLLEGE OF INFORMATION AND COMMUNICATION TECHNOLOGIES", level=2
    )
    college_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    department = doc.add_heading(f"DEPARTMENT OF {department}", level=2)
    department.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Practical Training Log - Book
    pt_title = doc.add_heading("PRACTICAL TRAINING LOG – BOOK", level=2)
    pt_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    # Student Information
    student_table = doc.add_table(rows=1, cols=2)
    student_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set column widths (50% for each column)
    student_table.columns[0].width = Inches(4)
    student_table.columns[1].width = Inches(2)

    # Add text to the cells
    cell_0_0 = student_table.cell(0, 0)
    cell_0_1 = student_table.cell(0, 1)

    cell_0_0.text = f"STUDENT NAME: {student_name}"
    cell_0_1.text = f"REG. NO: {reg_no}"

    # add another row with a single column for company name
    company_row = student_table.add_row()
    company_cell = company_row.cells[0]
    company_cell.text = f"COMPANY/INSTITUTION: {company}"
    company_cell.merge(student_table.cell(1, 1))
    doc.add_paragraph()

    week_log_table = doc.add_table(rows=1, cols=3)
    week_log_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Week number, from date, and to date cells
    week_number_cell = week_log_table.cell(0, 0)
    week_number_cell.text = f"WEEK NO: {week_no}"
    week_number_cell.paragraphs[0].paragraph_format.space_before = Inches(0.1)

    from_date_cell = week_log_table.cell(0, 1)
    from_date_cell.text = f"FROM: {from_date}"
    from_date_cell.paragraphs[0].paragraph_format.space_before = Inches(0.1)

    to_date_cell = week_log_table.cell(0, 2)
    to_date_cell.text = f"TO: {to_date}"
    to_date_cell.paragraphs[0].paragraph_format.space_before = Inches(0.1)

    # set cell widths 30% each
    week_log_table.columns[0].width = Inches(1)
    week_log_table.columns[1].width = Inches(2.5)
    week_log_table.columns[2].width = Inches(2.5)

    # Add borders to the table cells
    for cell in week_log_table._cells:
        set_cell_borders(cell)

    # add an empty paragraph with 2 new line breaks
    doc.add_paragraph()

    # Days and Activities
    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]

    # Create a table for Days and Activities
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False

    # Set vertical alignment to top for all cells
    for cell in table._cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Add borders to the table cells
        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)

    # Set column widths
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(4)

    # Add Days and Activities column headers
    cell_0_0 = table.cell(0, 0)
    cell_0_1 = table.cell(0, 1)

    cell_0_0.text = "DAY / DATE"
    cell_0_1.text = "ACTIVITY"

    # Make the text in the cells bold
    for cell in [cell_0_0, cell_0_1]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Align the text in the cells at the center
    cell_0_0.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_0_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add spacing before the headers
    cell_0_0.paragraphs[0].paragraph_format.space_before = Inches(0.1)
    cell_0_1.paragraphs[0].paragraph_format.space_before = Inches(0.1)

    # Populate the table with Days, Dates, and an empty cell for activities
    for i, day in enumerate(days):
        table.cell(i + 1, 0).text = f" {day} \n \n {data_dictionary[day]['date']}"
        table.cell(i + 1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.cell(i + 1, 1).text = data_dictionary[day]["activity"]

    doc.add_paragraph()

    # Details of the Main Job of the Week
    doc.add_heading("Details Of the Main Job of the Week", level=2)

    # Determine the number of rows in the machinery table based on the number of items in operations list
    machinery_rows = max(2, len(operations) + 1)  # Ensure at least 1 row for the header

    # Create the machinery table
    machinery_table = doc.add_table(
        rows=machinery_rows, cols=2
    )  # Add 1 row for the header
    machinery_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    machinery_table.autofit = False
    machinery_table.columns[0].width = Inches(3)
    machinery_table.columns[1].width = Inches(3)

    # Add borders to the table cells
    for cell in machinery_table._cells:
        set_cell_borders(cell)

    # Add the column headers (single-row header)
    cell_0_0 = machinery_table.cell(0, 0)
    cell_0_1 = machinery_table.cell(0, 1)

    cell_0_0.text = "Operation: "
    cell_0_1.text = "Machinery/Tools Used"

    # Make the text in the header cells bold
    for cell in [cell_0_0, cell_0_1]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Align the text in the header cells at the center
    cell_0_0.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell_0_1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Populate the machinery table with operations data
    for i, op_data in enumerate(operations, 1):
        op_cell = machinery_table.cell(i, 0)
        machinery_cell = machinery_table.cell(i, 1)

        op_cell.text = op_data["operation"]
        machinery_cell.text = op_data["machinery"]

    doc.add_paragraph()
    # table for comments from industrial supervisor, 1 row, 2 columns comments and signature
    comments_table = doc.add_table(rows=1, cols=2)
    comments_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    comments_table.columns[0].width = Inches(2)
    comments_table.columns[1].width = Inches(4)

    # Add borders to the table cells
    for cell in comments_table._cells:
        set_cell_borders(cell)

    # Add the column headers
    comments_table.cell(0, 0).text = "Comments from Industrial Supervisor"
    comments_table.cell(0, 1).text = "\n"

    # table for Name and Signature
    name_table = doc.add_table(rows=1, cols=2)
    name_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_table.columns[0].width = Inches(3)
    name_table.columns[1].width = Inches(3)

    # Add borders to the table cells
    for cell in name_table._cells:
        set_cell_borders(cell)

    # Add the column headers
    name_table.cell(
        0, 0
    ).text = "\nName: ............................................................"
    name_table.cell(0, 1).text = "\nSignature: ......................................."

    doc.add_page_break()  # page break

    # Detailed Diagram of the Main Job
    diagram_paragraph = doc.add_paragraph()
    diagram_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # table with 3 rows and 4 columns for the diagram
    # first row is for the title 'Detailed Diagram of the Main Job' aligned center
    # second row is for the diagram image with all 4 columns merged
    # third row is for the 'drawn by', 'date', 'checked by', 'date' aligned left
    diagram_table = doc.add_table(rows=3, cols=4)
    diagram_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    diagram_table.autofit = False
    diagram_table.columns[0].width = Inches(1.5)
    diagram_table.columns[1].width = Inches(1.5)
    diagram_table.columns[2].width = Inches(1.5)
    diagram_table.columns[3].width = Inches(1.5)

    # Add borders to the table cells
    for cell in diagram_table._cells:
        set_cell_borders(cell)

    # Add the column headers
    diagram_table.cell(0, 0).text = "Detailed Diagram of the Main Job"
    diagram_table.cell(0, 0).merge(diagram_table.cell(0, 3))
    diagram_table.cell(1, 0).merge(diagram_table.cell(1, 3))
    diagram_table.cell(
        2, 0
    ).text = "Drawn by: \n\n........................................"
    diagram_table.cell(2, 1).text = "Date: \n\n........................................"
    diagram_table.cell(
        2, 2
    ).text = "Checked by: \n\n........................................"
    diagram_table.cell(2, 3).text = "Date: \n\n........................................"

    # Add the diagram image
    diagram_table.cell(1, 0).merge(diagram_table.cell(1, 3))
    diagram_table.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # get image from logbook.activity_diagram
    if activity_diagram and os.path.exists(activity_diagram):
        diagram_table.cell(1, 0).paragraphs[0].add_run().add_picture(
            activity_diagram, width=Inches(6)
        )
    else:
        # Add placeholder text if no diagram is available
        diagram_table.cell(1, 0).text = "No diagram available"

    # Save the Document
    # Ensure the docs directory exists
    docs_dir = os.path.join(MEDIA_ROOT, 'docs')
    os.makedirs(docs_dir, exist_ok=True)
    
    filepath = os.path.join(
        docs_dir, f"{reg_no}-week-{week_no}-practical-training-logbook.docx"
    )
    doc.save(filepath)

    return filepath
